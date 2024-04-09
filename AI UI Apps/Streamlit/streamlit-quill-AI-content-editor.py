import os
import openai
import streamlit as st
import assemblyai as aai
import streamlit_quill

from streamlit_chat import message
from streamlit_quill import st_quill
from streamlit_modal import Modal

#text editor export
import docx
from docx import Document
from PyPDF2 import PdfReader
import base64
from io import BytesIO
import embed_pdf

#load environment variables
from dotenv import load_dotenv
load_dotenv("../.env")

#Azure OpenAI Imports
from openai import AzureOpenAI
from langchain_openai.chat_models import AzureChatOpenAI

from llama_index import SimpleDirectoryReader, VectorStoreIndex

#Azure Auth
AZURE_OPENAI_API_KEY = os.getenv('AZURE_OPENAI_API_KEY')
AZURE_OPENAI_ENDPOINT = os.getenv('AZURE_OPENAI_ENDPOINT')
AZURE_OPENAI_DEPLOYMENT = os.getenv('AZURE_OPENAI_DEPLOYMENT')
AZURE_OPENAI_EMBEDDING_DEPLOYMENT = os.getenv('AZURE_OPENAI_EMBEDDING_DEPLOYMENT')
AZURE_OPENAI_API_VERSION = os.getenv('AZURE_OPENAI_API_VERSION')

#RAG Features
from langchain_community.vectorstores.faiss import FAISS
import pandas as pd
file_formats = {
        "csv": pd.read_csv,
        "xls": pd.read_excel,
        "xlsx": pd.read_excel,
        "xlsm": pd.read_excel,
        "xlsb": pd.read_excel,
    }

# Function to convert text to DOCX
def text_to_docx(text):
    doc = Document()
    doc.add_paragraph(text)
    return doc

# Function to download the DOCX file
def download_docx(doc, filename="exported_content.docx"):
    with BytesIO() as buffer:
        doc.save(buffer)
        buffer.seek(0)
        b64 = base64.b64encode(buffer.read()).decode()
    href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}">Download as DOCX</a>'
    return href

# Define constants and settings
CACHE_DIR = "./uploads"
aai.settings.api_key = st.secrets['assembly_api_key']

# Cache the result to avoid recomputation
@st.cache_resource(show_spinner="Indexing documents...Please have patience")
def build_index(files):
    documents = SimpleDirectoryReader(input_files=files).load_data()
    index = VectorStoreIndex.from_documents(documents)
    return index

# Transcribe audio and video files
def transcribe_audio_video(file_path):
    transcriber = aai.Transcriber()
    transcript = transcriber.transcribe(file_path)
    transcript_path = file_path + ".txt"
    with open(transcript_path, "w") as f:
        f.write(transcript.text)
    return transcript_path

# Upload files and cache them
def upload_files(types=["pdf", "txt", "mp3", "mp4", 'mpeg', 'doc', 'docx'], **kwargs):
    files = st.file_uploader(
        label=f"Upload files", type=types, **kwargs
    )
    if not files:
        st.info(f"No documents uploaded")
        #st.stop()
    return cache_files(files, types=types)

# Cache uploaded files
def cache_files(files, types=["pdf", "txt", "mp3", "mp4", 'mpeg', 'doc', 'docx']) -> list[str]:
    filepaths = []
    for file in files:
        # Determine the file extension from the mime type
        ext = file.type.split("/")[-1]
        if ext == "plain":  # Handle text/plain mime type
            ext = "txt"
        elif ext in ["vnd.openxmlformats-officedocument.wordprocessingml.document", "vnd.ms-word"]:
            ext = "docx"  # or "doc" depending on your needs
        if ext not in types:
            continue
        filepath = f"{CACHE_DIR}/{file.name}"
        with open(filepath, "wb") as f:
            f.write(file.getvalue())
        if ext in ["mp3", "mp4"]:
            filepath = transcribe_audio_video(filepath)
        filepaths.append(filepath)
    # st.sidebar.write("Uploaded files", filepaths)  # Debug statement
    with st.sidebar:
        with st.expander("Uploaded Files"):
            filepaths_pretty = "\n".join(f"- {filepath}" for filepath in filepaths)
            st.markdown(f"{filepaths_pretty}")
    return filepaths

def transcribe_and_save(file_path):
    transcriber = aai.Transcriber()
    transcript = transcriber.transcribe(file_path)
    transcript_path = file_path + ".txt"
    with open(transcript_path, "w") as f:
        f.write(transcript.text)
    return transcript_path

# Save extracted text to a txt file
def save_extracted_text_to_txt(text, filename):
    txt_filename = os.path.splitext(filename)[0] + ".txt"
    txt_filepath = os.path.join('uploads_txt', txt_filename)
    with open(txt_filepath, 'w', encoding='utf-8') as txt_file:
        txt_file.write(text)
    return txt_filepath

# Get OpenAI API key from session state
def get_key():
    return st.session_state["openai_api_key"]

# Read text from Word document
def read_word_file(file_path):
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

# Process PDF files
def process_pdf(pdf_path):
    reader = PdfReader(pdf_path)
    all_text = ""
    for page in reader.pages:
        extracted_text = page.extract_text()
        if extracted_text:
            processed_text = ' '.join(extracted_text.split('\n'))
            all_text += processed_text + "\n\n"
    txt_filepath = save_extracted_text_to_txt(all_text, os.path.basename(pdf_path))
    #os.remove(pdf_path)  # Delete the original PDF file
    return txt_filepath

# Process uploaded documents
def process_documents(documents):
    processed_docs = []
    for doc in documents:
        if doc.endswith('.pdf'):
            processed_docs.append(process_pdf(doc))
        elif doc.endswith(('.doc', '.docx')):
            text = read_word_file(doc)
            txt_filepath = save_extracted_text_to_txt(text, os.path.basename(doc))
            processed_docs.append(txt_filepath)
        elif doc.endswith(('.mp3', '.mp4', '.mpeg')):
            processed_docs.append(transcribe_and_save(doc))
        else:
            processed_docs.append(doc)
    return processed_docs



# Setting page title and header
st.set_page_config(page_title="TAC | TCD AI Companion", page_icon=":brain:", layout="wide",initial_sidebar_state="expanded")
st.markdown("<h1 style='text-align: center;'>Backlog Content Studio 🎨</h1>", unsafe_allow_html=True)
st.divider()

# Initialise session state variables
if 'generated' not in st.session_state:
    st.session_state['generated'] = []
if 'past' not in st.session_state:
    st.session_state['past'] = []
if 'messages' not in st.session_state:
    st.session_state['messages'] = [
        {"role": "system", "content": "You are an expert scrum master who is proficient at creating any type of backlog content, such as user stories, epics, and bugs. When you are asked to create backlog content you create it in a way that is clear and properly formatted. If it's a user story it has clear acceptance criteria and follows the standards of well crafted user stories. when you respond to the user if they ask for some type of backlog content, be sure to respond using an easily readable structured numbered outline format."}
    ]
if 'model_name' not in st.session_state:
    st.session_state['model_name'] = []
if 'cost' not in st.session_state:
    st.session_state['cost'] = []
if 'total_tokens' not in st.session_state:
    st.session_state['total_tokens'] = []
if 'total_cost' not in st.session_state:
    st.session_state['total_cost'] = 0.0

client = AzureOpenAI(
    azure_deployment= AZURE_OPENAI_DEPLOYMENT,
    api_version=AZURE_OPENAI_API_VERSION,
    api_key = AZURE_OPENAI_API_KEY ,
    azure_endpoint = AZURE_OPENAI_ENDPOINT,
    #streaming = True,
    #temperature=0
    # is_chat_model=False,  # Comment in if you deployed a completion model
)

# Sidebar - let user choose model, show total cost of current conversation, and let user clear the current conversation
with st.sidebar:
    st.sidebar.title("Sidebar")
    # Displaying information and warnings in the sidebar
    st.sidebar.info(
        "Do not submit information such as personally identifiable information, client information, or confidential company information."
    )
    st.sidebar.warning(
        "LLMs may produce inaccurate information about people, places, or facts. Don't entirely trust them."
    )
    SKILL = st.selectbox(label='Choose Skill', options=['Backlog Content Studio','Agile Operating Model','Agile Coach','Analytics & Reporting','Scaled Agile, Portfolio/Program', 'Planning Assistant','Document Analysis'])
    with st.expander("Advanced Settings", expanded=False):        
        model_name = st.selectbox("Choose a model:", ("GPT-4 Turbo", "GPT-3.5 Turbo"))
        TEMP = st.slider("Temperature",0.0,1.0,0.1)

    modal = Modal("Backlog Content Wizard :male_mage:", key="user_input")
    if st.button('Content Wizard :male_mage:'):
        modal.open()

    if modal.is_open():
        with modal.container():
            tab1, tab2 = st.tabs(["Wizard","Options"])

            with tab1:                
                with st.form('wizard'):
                    typeselection = st.selectbox(
                        "Select type of content:",
                        ("User Stories", "Epics", "Spikes", "User Journey", "User Acceptance Tests", "Non-functional Requirements"))
                    quantity = st.number_input("How many backlog items do you want to create?", value=5, placeholder="Type number here...", min_value=1, max_value=20)
                    personaq1 = st.text_area("Who is the specific audience for the content or type of user of the feature? (Business Stakholder, End-User, Developer, Product Owner, etc)", placeholder="Type answer here...")
                    personaq2 = st.text_area("what are the user pain points, opportunities or benefits this feature intends to address?", placeholder="Type answer here...")
                    personaq3 = st.text_area("How will the persona or user use the feature? How will the content or feature help the persona be successful?", placeholder="Type answer here...")
                    question1 = st.text_area("Which unique value does this feature offer to the users?", placeholder="Type answer here...")
                    question2 = st.text_area("Which specific aspect or functionality of the product does this content primarily aim to create or enhance?", placeholder="Type answer here...") 
                    question3 = st.text_area("What is the is the central goal or hypothesis that this feature aims to validate or achieve?", placeholder="Type answer here...")
                    question4 = st.text_area("What technology considerations or technical constraints must this feature accomodate or meet?", placeholder="Type answer here...")
                    question5 = st.text_area("What is the step by step process that this feature must follow or the user must follow?", placeholder="Type answer here...")
                    submit = st.form_submit_button('Submit')
                
            with tab2:
                st.header("Options")
                typeoptions = st.selectbox(
                    "Select type of User Story:",
                    ("Standard (Default)", "Behavioral Drive Development (BDD)", "Hypothesis Drive Development (HDD)", "Test Driven Development (TDD)", "Gherkin"))
                if typeoptions == "Standard (Default)":
                    st.write("As a User, I want to... So that...")
                elif typeoptions == "Behavioral Drive Development (BDD)":
                    st.write(f"""As a User, I want to... So that... Scenario 1: Given... when... then...\n
                    As a Backlog content studio user I want to create 10 user stories so that I can effectively communicate functionality to developers and cut time off the process.\n
                    Scenario 1: Given I'm on the homepage
                    """)  
                elif typeoptions == "Hypothesis Drive Development (HDD)":
                    st.write("You selected the Hypothesis Driven Development user story type")
                elif typeoptions == "Test Driven Development (TDD)":
                    st.write("You selected the Test Driven Development user story type")
                else:
                    st.write("You selected the Gherkin user story type")
                
                ACoptions = st.selectbox(
                    "Select type of Acceptance Criteria:",
                    ("Standard (Default)", "Scenario Based", "Behavioral Drive Development (BDD)","Hypothesis Drive Development (HDD)", "Test Driven Development (TDD)", "Gherkin"))

                submitcol1, cancelcol2 = st.columns(2)

                with submitcol1:

                    if st.button("Submit", key="submit question"):
                        message = f"Q1: {question1}\nQ2: {question2}\nQ3: {question3}\nQ4: {question4}\nQ5: {question5}\nQ6: {quantity}"
                        response = client.chat.completions.create(prompt=message, model="text-davinci-003")
                        st.write(response.choices[0].text)
                        modal.close()

                with cancelcol2:
                    if st.button("Cancel"):
                            modal.close()

    st.divider()

    counter_placeholder = st.sidebar.empty()
    clear_button = st.sidebar.button("Clear Conversation", key="clear")
    counter_placeholder.write(f"Total cost of this conversation: ${st.session_state['total_cost']:.5f}")

    st.divider()

    # Upload PDFs, DOCs, TXTs, MP3s, and MP4s
    documents = upload_files(types=["pdf", "txt", "mp3", "mp4", 'mpeg', 'doc', 'docx'], accept_multiple_files=True)

    # Process the uploaded documents
    processed_documents = process_documents(documents)

    # Check if there are any processed documents before building the index
    if processed_documents:
        index = build_index(processed_documents)
        if st.button("Embed Documents"):
            st.info("Embedding documents...")
            try:
                embed_pdf.embed_all_pdf_docs()
                st.sidebar.info("Done!")
            except Exception as e:
                st.sidebar.error(e)
                st.sidebar.error("Failed to embed documents.")
    else:
        # Handle the case where no documents are uploaded/processed
        # For example, you could set `index` to None or provide a default value
        index = None
        #st.info("No documents uploaded. Please upload documents to proceed.")

    with st.expander("Advanced Settings", expanded=False):
        chain_type = st.selectbox("Chain Type", ["map_reduce", "stuff", "refine"])
        chunksize = st.number_input("Chunk Size:", value=200, placeholder="Type number here...", min_value=100)
        chunkOverlap = st.number_input("Chunk Overlap:", value=100, placeholder="Type number here...", min_value=50)
        limit = st.number_input("Retrieve X most relevant results:", value=6, placeholder="Type number here...", min_value=1, max_value=20)
        certainty = st.slider("Certainty threshold for relevancy", 0.0, 1.0, 0.75)

    st.divider()

# Map model names to OpenAI model IDs
if model_name == "GPT-3.5 Turbo":
    model = "jt-gpt4-turbo"
else:
    if model_name == "GPT-4 Turbo":
        model = "jt-gpt4-turbo"
    else:
        model = "jt-gpt4-turbo"

client = AzureOpenAI(
    azure_deployment=model,
    api_version=AZURE_OPENAI_API_VERSION,
    api_key = AZURE_OPENAI_API_KEY ,
    azure_endpoint = AZURE_OPENAI_ENDPOINT,
    #streaming = True,
    #temperature=0
    # is_chat_model=False,  # Comment in if you deployed a completion model
)

# client2 = AzureChatOpenAI(
#     azure_deployment=model,
#     api_version="2023-08-01-preview",
#     api_key = AZURE_OPENAI_API_KEY ,
#     azure_endpoint = "https://caedaoipocaoa0p.openai.azure.com/",
#     streaming = True,
#     temperature=0
#     # is_chat_model=False,  # Comment in if you deployed a completion model
# )

# reset everything
if clear_button:
    st.session_state['generated'] = []
    st.session_state['past'] = []
    st.session_state['messages'] = [
        {"role": "system", "content": "You are a helpful assistant."}
    ]
    st.session_state['number_tokens'] = []
    st.session_state['model_name'] = []
    st.session_state['cost'] = []
    st.session_state['total_cost'] = 0.0
    st.session_state['total_tokens'] = []
    counter_placeholder.write(f"Total cost of this conversation: ${st.session_state['total_cost']:.5f}")

# generate a response
def generate_response(prompt):
    st.session_state['messages'].append({"role": "user", "content": prompt})

    completion = client.chat.completions.create(
        model=model,
        messages=st.session_state['messages']
    )
    response = completion.choices[0].message.content
    st.session_state['messages'].append({"role": "assistant", "content": response})

    # print(st.session_state['messages'])
    total_tokens = completion.usage.total_tokens
    prompt_tokens = completion.usage.prompt_tokens
    completion_tokens = completion.usage.completion_tokens
    return response, total_tokens, prompt_tokens, completion_tokens

# Initialize a session state variable for the editor content if not already done
if 'editor_content' not in st.session_state:
    st.session_state['editor_content'] = ""

# Function to update the editor content, with an option to append
def update_editor_content(new_content, append=False):
    if append:
        # Append the new content to the existing content in the editor
        st.session_state['editor_content'] += "\n\n" + new_content
    else:
        # Replace the editor's content with the new content
        st.session_state['editor_content'] = new_content

col1, col2 = st.columns(2)

with col1:
    st.title("Content Companion 🧙‍♂️")
    # container for chat history
    response_container = st.container()
    # container for text box
    container = st.container()

    with container:
        with st.form(key='my_form', clear_on_submit=True):
            user_input = st.text_area("You:", key='input', height=100)
            submit_button = st.form_submit_button(label='Send')

        if submit_button and user_input:
            with st.spinner('Generating response...'):
                output, total_tokens, prompt_tokens, completion_tokens = generate_response(user_input)
            st.session_state['past'].append(user_input)
            st.session_state['generated'].append(output)
            st.session_state['model_name'].append(model_name)
            st.session_state['total_tokens'].append(total_tokens)

            # from https://openai.com/pricing#language-models
            if model_name == "GPT-3.5":
                cost = total_tokens * 0.0015 / 1000
            else:
                cost = (prompt_tokens * 0.01 + completion_tokens * 0.03) / 1000

            st.session_state['cost'].append(cost)
            st.session_state['total_cost'] += cost

    if st.session_state['generated']:
        with response_container:
            for i in range(len(st.session_state['generated'])):
                message(st.session_state["past"][i], is_user=True, key=str(i) + '_user')
                message(st.session_state["generated"][i], key=str(i))
                # Existing "Copy to Editor" button
                if st.button(label="Copy to Editor", key=f"copy_btn_{i}"):
                    update_editor_content(st.session_state["generated"][i], append=False)
                # New "Append to Editor" button
                if st.button(label="Append to Editor", key=f"append_btn_{i}"):
                    update_editor_content(st.session_state["generated"][i], append=True)
                st.write(
                    f"Model used: {st.session_state['model_name'][i]}; Number of tokens: {st.session_state['total_tokens'][i]}; Cost: ${st.session_state['cost'][i]:.5f}")
                counter_placeholder.write(f"Total cost of this conversation: ${st.session_state['total_cost']:.5f}")

with col2:
    st.title("Content Editor 📝")
    # Render the Quill editor with the content from the session state
    editor_content = st_quill(value=st.session_state.get('editor_content', ""))
    
    # Optionally, add a button to clear the editor
    if st.button("Clear Editor"):
        update_editor_content("")

    # Export to DOCX button
    if st.button("Export to DOCX"):
        doc = text_to_docx(st.session_state.get('editor_content', ""))
        st.markdown(download_docx(doc), unsafe_allow_html=True)
